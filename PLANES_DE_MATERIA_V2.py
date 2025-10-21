"""
PLANES_DE_MATERIA_V2
"""

#!/usr/bin/env python
# coding: utf-8

# In[1]:


import os
from openai import OpenAI
from docx import Document
from dotenv import load_dotenv

# Cargar variables de entorno desde el archivo .env
load_dotenv()

# Obtener clave API de OpenAI desde variable de entorno
api_key = os.getenv("OPENAI_API_KEY")

if not api_key:
    raise ValueError("No se encontró la clave de API en el archivo .env")

client = OpenAI(api_key=api_key)

# Función para crear el asistente
def create_assistant():
    try:
        assistant = client.beta.assistants.create(
            name="Asistente de creación de planes de materia",
            instructions="""Eres un experto pedagogo en crear planes de materia,
            bajo la corriente de crear competencias en cada estudiante. 
            No proporcionas las fuentes de información.""",
            model="gpt-4o",
            tools=[{"type": "file_search"}],
        )
        return assistant
    except Exception as e:
        print(f"Error creating assistant: {e}")
        return None

# Función para subir archivos al vector store
def upload_files_to_vector_store(directory, vector_store):
    try:
        file_paths = [
            os.path.join(directory, file) for file in os.listdir(directory) if os.path.isfile(os.path.join(directory, file))
        ]
        file_streams = [open(path, "rb") for path in file_paths]

        # Subir archivos y monitorear el progreso
        file_batch = client.beta.vector_stores.file_batches.upload_and_poll(
            vector_store_id=vector_store.id, files=file_streams
        )

        print(f"Estado de subida: {file_batch.status}")
        print(f"Cantidad de archivos subidos: {file_batch.file_counts}")
    except Exception as e:
        print(f"Error al subir archivos: {e}")
    finally:
        for file_stream in file_streams:
            file_stream.close()

# Función para crear el plan de materia
def create_plan_for_subject(assistant, subject_name, lessons, vector_store):
    try:
        assistant = client.beta.assistants.update(
            assistant_id=assistant.id,
            tool_resources={"file_search": {"vector_store_ids": [vector_store.id]}},
        )

        thread = client.beta.threads.create(
            messages=[
                {
                    "role": "user",
                    "content": f"""Propuesta del plan de materia para {subject_name}:
                    1. JUSTIFICACIÓN DE LA ASIGNATURA (500-1000 tokens)
                    2. COMPETENCIA PROPIA DE LA ASIGNATURA
                    3. COMPETENCIA GENÉRICA O TRANSVERSAL
                    4. Competencias específicas obtenidas (4 principales)
                    5. Lista de ELEMENTOS DE COMPETENCIA en línea con la competencia genérica
                    6. Títulos de lección o unidad para:
                    - {lessons[0]}
                    - {lessons[1]}
                    - {lessons[2]}
                    - {lessons[3]}
                    - {lessons[4]}
                    7. Actividades didácticas para cada lección
                    8. SABERES PROCEDIMENTALES para cada elemento
                    9. SABERES CONCEPTUALES para cada elemento
                    10. SABERES ACTITUDINALES para cada elemento
                    11. PERFIL PROFESIONAL del estudiante al finalizar el curso
                    """
                }
            ]
        )

        print(f"Hilo creado para {subject_name}.")
        return thread
    except Exception as e:
        print(f"Error creando el hilo: {e}")
        return None

# Función para ejecutar el asistente y obtener la respuesta
def run_assistant_and_get_response(thread, assistant, directory_out):
    try:
        run = client.beta.threads.runs.create_and_poll(
            thread_id=thread.id, assistant_id=assistant.id
        )

        # Obtener los mensajes del resultado
        messages = list(client.beta.threads.messages.list(thread_id=thread.id, run_id=run.id))
        print(messages[0].content[0].text.value)
        message_content = messages[0].content[0].text.value

        save_response_to_docx(message_content, directory_out)
    except Exception as e:
        print(f"Error ejecutando el asistente: {e}")


# Función para guardar la respuesta en un archivo DOCX
def save_response_to_docx(content, file_path):
    try:
        document = Document()
        document.add_heading('Plan de Materia', level=1)
        document.add_paragraph(content)
        document.save(file_path)
        print(f"Respuesta guardada en {file_path}")
    except Exception as e:
        print(f"Error al guardar el archivo DOCX: {e}")

# Función principal
def main():
    subject_name = "Introducción al Excel"
    directory = "C:\\Users\\HP\\Desktop\\Planificación Curricular\\Administración\\"
    directory_out = f"C:\\Users\\HP\\Desktop\\PROPUESTAS DE DIPLOMADO\\PROPUESTAS DE DIPLOMADO A POS-GRADO\\PROPUESTAS\DIPLOMADO EXCEL Y POWER BOI APLICADO A LA GESTION DE VENTAS\\PLANES DE MATERIA\\{subject_name}.docx"

    lessons = [
                "Fundamentos de excel",
                "Manejo de datos en excel",
                "Análisis de datos con funciones y fórmulas",
                "Visualización de datos con gráficos",
                "Tablas dinámicas y análisis"
                ]

    assistant = create_assistant()
    if not assistant:
        return

    vector_store = client.beta.vector_stores.create(name=f"Planes de materia - {subject_name}")
    upload_files_to_vector_store(directory, vector_store)

    thread = create_plan_for_subject(assistant, subject_name, lessons, vector_store)
    if thread:
        run_assistant_and_get_response(thread, assistant, directory_out)

if __name__ == "__main__":
    main()


# In[ ]:





# In[ ]:




